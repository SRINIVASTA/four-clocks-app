import streamlit as st
import pandas as pd
import re
import io
from datetime import datetime

# Document conversion imports
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

# 1. PAGE LAYOUT SETUP
st.set_page_config(
    page_title="⏰ tasks-four-clocks-app",
    page_icon="⏰",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Initialize Session State tracking database arrays with upgraded columns
if 'tasks_db' not in st.session_state:
    schema = {
        'task_title': pd.Series(dtype='object'),
        'clock_bucket': pd.Series(dtype='category'),
        'status': pd.Series(dtype='string'), # Changed to string for fluid state management
        'is_critical': pd.Series(dtype='bool'), # New priority data flag
        'date_created': pd.Series(dtype='datetime64[ns]')
    }
    st.session_state.tasks_db = pd.DataFrame(schema)

# 2. KEYWORD ROUTING LAWS ENGINE
def predict_clock_bucket(task_text):
    text_lower = task_text.lower()
    if any(w in text_lower for w in ['sql', 'exploratory data', 'underwriting engine', 'asset records', 
                                    'gpu', 'api', 'docker', 'bug', 'client', 'tax', 'audit', 'security',
                                    'cache', 'redis', 'pipeline', 'kafka', 'database', 'mongodb', 'postgresql',
                                    'fastapi', 'penetration', 'server', 'receipts', 'accounting', 'invoice']):
        return 'Now Clock ⏱️'
    elif any(w in text_lower for w in ['linkedin', 'kaggle', 'newsletter', 'submission', 'youtube', 
                                      'scripts', 'outreach', 'pricing page', 'scrape', 'podcast', 
                                      'reddit', 'emails', 'marketing', 'draw', 'giveaways']):
        return 'Compound Clock 📈'
    elif any(w in text_lower for w in ['document parser', 'rag and llm', 'portfolio risk', 'training modules', 
                                      'whitepaper', 'legal', 'nondisclosure', 'research paper', 'compliance', 
                                      'governance', 'patent', 'roadmap', 'agreement', 'addendum']):
        return 'Deep Clock 🧠'
    else:
        return 'Wild Clock ⚡'

def split_text_into_clean_tasks(raw_text_block):
    sentences = re.split(r'\.|\n', raw_text_block)
    processed_list = []
    for s in sentences:
        clean = s.strip()
        if len(clean) > 5:
            processed_list.append(clean + ".")
    return processed_list

def add_tasks_to_database(task_list, marks_critical=False):
    new_rows = []
    for t in task_list:
        # Check duplicate pipeline rule
        if not st.session_state.tasks_db.empty and t.lower() in st.session_state.tasks_db['task_title'].str.lower().values:
            continue
        assigned_bucket = predict_clock_bucket(t)
        new_rows.append({
            'task_title': t,
            'clock_bucket': assigned_bucket,
            'status': '⏳ Pending',
            'is_critical': marks_critical,
            'date_created': datetime.now()
        })
    if new_rows:
        new_df = pd.DataFrame(new_rows)
        st.session_state.tasks_db = pd.concat([st.session_state.tasks_db, new_df], ignore_index=True)

# 3. FILE PARSING PIPELINE
def extract_tasks_from_file(uploaded_file):
    ext = uploaded_file.name.split('.')[-1].lower()
    text_data = ""
    try:
        if ext == 'txt':
            text_data = uploaded_file.read().decode('utf-8')
        elif ext == 'csv':
            df = pd.read_csv(uploaded_file)
            text_data = " ".join(df.astype(str).values.flatten())
        elif ext == 'xlsx':
            df = pd.read_excel(uploaded_file)
            text_data = " ".join(df.astype(str).values.flatten())
        elif ext == 'docx':
            doc = Document(io.BytesIO(uploaded_file.read()))
            text_data = " ".join([p.text for p in doc.paragraphs])
        return split_text_into_clean_tasks(text_data)
    except Exception as e:
        st.error(f"Error parsing file metadata stream: {e}")
        return []

# 4. REPORT EXPORT BINARY GENERATORS
def generate_export_bytes(format_type):
    df = st.session_state.tasks_db
    if df.empty:
        return None
    buffer = io.BytesIO()
    
    if format_type == "CSV":
        # Create a temporary copy to protect the live database state
        export_df = df.copy()
        
        # Convert the datetime column into a clean string format Excel understands
        if 'date_created' in export_df.columns and not export_df.empty:
            export_df['date_created'] = export_df['date_created'].dt.strftime('%Y-%m-%d %H:%M:%S')
            
        export_df.to_csv(buffer, index=False, encoding='utf-8-sig')
        return buffer.getvalue()
    elif format_type == "Excel":
        df.to_excel(buffer, index=False)
        return buffer.getvalue()
    elif format_type == "Text":
        text_out = "⏰ THE FOUR CLOCKS SPRINT REPORT\n" + "="*40 + "\n\n"
        for cat in ['Now Clock ⏱️', 'Compound Clock 📈', 'Deep Clock 🧠', 'Wild Clock ⚡']:
            text_out += f"=== {cat.upper()} ===\n"
            sub = df[df['clock_bucket'] == cat]
            for idx, row in sub.iterrows():
                crit_tag = "[CRITICAL] " if row['is_critical'] else ""
                text_out += f"- {crit_tag}{row['task_title']} ({row['status']})\n"
            text_out += "\n"
        return text_out.encode('utf-8')
    elif format_type == "Word":
        doc = Document()
        doc.add_heading('The Four Clocks Balanced Sprint', 0)
        for cat in ['Now Clock ⏱️', 'Compound Clock 📈', 'Deep Clock 🧠', 'Wild Clock ⚡']:
            doc.add_heading(cat, level=1)
            sub = df[df['clock_bucket'] == cat]
            for idx, row in sub.iterrows():
                prefix = "🚨 CRITICAL: " if row['is_critical'] else ""
                doc.add_paragraph(f"{prefix}{row['task_title']} ({row['status']})", style='List Bullet')
        doc.save(buffer)
        return buffer.getvalue()
    elif format_type == "PDF":
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        title_style = ParagraphStyle('T', parent=styles['Heading1'], fontSize=18, spaceAfter=20)
        cat_style = ParagraphStyle('C', parent=styles['Heading2'], fontSize=12, spaceBefore=12, spaceAfter=8)
        task_style = ParagraphStyle('N', parent=styles['Normal'], fontSize=10, leftIndent=15, spaceAfter=4)
        
        story.append(Paragraph("<b>The Four Clocks Dashboard Report</b>", title_style))
        for cat in ['Now Clock ⏱️', 'Compound Clock 📈', 'Deep Clock 🧠', 'Wild Clock ⚡']:
            story.append(Paragraph(f"<b>{cat.upper()}</b>", cat_style))
            sub = df[df['clock_bucket'] == cat]
            for idx, row in sub.iterrows():
                prefix = "<font color='red'><b>[🚨 CRITICAL]</b></font> " if row['is_critical'] else ""
                story.append(Paragraph(f"{prefix}{row['task_title']} ({row['status']})", task_style))
        doc.build(story)
        return buffer.getvalue()

# 5. DRAW GRAPHICAL FRONTEND LAYOUT INTERFACE
st.title("⏰ The Four Clocks Smart Workspace")
st.markdown("Automate your 15-minute Monday ritual. Developed by **Srinivasta**.")

# SIDEBAR INPUT PANEL MANAGEMENT
with st.sidebar:
    st.header("📥 Ingest Tasks Data")
    
    # Method 1: Text Box Sandbox Workspace
    st.subheader("Method 1: Paste Text Block")
    text_input = st.text_area("Paste mixed paragraphs here:", placeholder="Type or paste unstructured task entries...", height=120)
    
    # NEW: Toggle to inject critical status priority on entry batch
    mark_as_critical = st.checkbox("🔥 Mark this batch as High Priority (Critical)")
    
    if st.button("Process Text Block", type="primary", use_container_width=True):
        if text_input:
            tasks = split_text_into_clean_tasks(text_input)
            add_tasks_to_database(tasks, marks_critical=mark_as_critical)
            st.toast("⚡ Text block loaded successfully!", icon="✅")
            st.rerun()

    st.markdown("---")
    
    # Method 2: Document Drag and Drop Upload Panel
    st.subheader("Method 2: Ingest File")
    uploaded_file = st.file_uploader("Upload a file:", type=['txt', 'csv', 'xlsx', 'docx'])
    if uploaded_file is not None:
        file_tasks = extract_tasks_from_file(uploaded_file)
        if file_tasks:
            add_tasks_to_database(file_tasks, marks_critical=False)
            st.toast(f"📊 Extracted items from {uploaded_file.name}!", icon="📥")
            st.rerun()

    st.markdown("---")
    st.button("🗑️ Reset Application Database", type="secondary", on_click=lambda: st.session_state.clear(), use_container_width=True)

# MAIN COLUMN GRID WORKSPACE DISPLAY PANEL
if st.session_state.tasks_db.empty:
    st.info("💡 Your workspace dashboard is empty. Paste a running list or drop a file in the sidebar to populate data views.")
else:
    # 4-Column Horizon Matrix View
    col1, col2, col3, col4 = st.columns(4)
    clocks_meta = [
        ("Now Clock ⏱️", col1),
        ("Compound Clock 📈", col2),
        ("Deep Clock 🧠", col3),
        ("Wild Clock ⚡", col4)
    ]
    
    for bucket_name, col_obj in clocks_meta:
        with col_obj:
            # Filter and AUTOMATICALLY SORT: Newest items appear at the very top (date_created descending)
            filtered_df = st.session_state.tasks_db[st.session_state.tasks_db['clock_bucket'] == bucket_name]
            filtered_df = filtered_df.sort_values(by='date_created', ascending=False)
            
            st.subheader(bucket_name)
            st.metric(label="Total Items", value=len(filtered_df))
            
            if filtered_df.empty:
                st.caption("No matching task entries found.")
            else:
                for idx, row in filtered_df.iterrows():
                    # Create unique graphic borders for each task item module
                    with st.container(border=True):
                        # Feature 2: High Priority Red text Banner Trigger
                        if row['is_critical']:
                            st.markdown("<span style='color:#FF4B4B; font-weight:bold; font-size:14px;'>🚨 CRITICAL PRIORITY</span>", unsafe_allow_html=True)
                        st.write(row['task_title'])
                        
                        # Feature 3: Status Pipeline Selector Hub
                        current_status = row['status']
                        status_options = ['⏳ Pending', '⚡ In Progress', '✅ Completed']
                        
                        # Find matching index to avoid rendering dropdown mismatches
                        try:
                            status_idx = status_options.index(current_status)
                        except ValueError:
                            status_idx = 0
                            
                        chosen_status = st.selectbox(
                            "Update Status:",
                            options=status_options,
                            index=status_idx,
                            key=f"status_{idx}",
                            label_visibility="collapsed"
                        )
                        
                        # Handle live pipeline dataframe mutation shifts
                        if chosen_status != current_status:
                            st.session_state.tasks_db.at[idx, 'status'] = chosen_status
                            st.rerun()

    # EXPORT DOWNSTREAM COMPILATION PANEL MANAGER
    st.markdown("---")
    st.markdown("### 📥 Compile Work Cycle Outputs")
    exp_col1, exp_col2 = st.columns([1, 3])
    with exp_col1:
        export_format = st.selectbox("Choose format:", ["Text", "PDF", "Word", "Excel", "CSV"], label_visibility="collapsed")
    with exp_col2:
        binary_data = generate_export_bytes(export_format)
        ext_map = {"Text": "txt", "PDF": "pdf", "Word": "docx", "Excel": "xlsx", "CSV": "csv"}
        if binary_data:
            st.download_button(
                label=f"Download {export_format} Document",
                data=binary_data,
                file_name=f"four_clocks_sprint.{ext_map[export_format]}",
                mime="application/octet-stream",
                type="primary"
            )
